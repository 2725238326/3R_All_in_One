# -*- coding: utf-8 -*-
"""重写结题报告全文：学术风格，基于真实项目数据。"""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1]
OUT = sys.argv[2]

doc = Document(SRC)
P = doc.paragraphs
CEN = WD_ALIGN_PARAGRAPH.CENTER

STYLE_BODY = P[47].style
STYLE_H1   = P[45].style
STYLE_H2   = P[46].style
STYLE_REF  = P[99].style

# ── helper ────────────────────────────────────────────────────────────
def set_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def insert_after(anchor_p, text, style, align=None):
    new_p = OxmlElement('w:p')
    anchor_p._p.addnext(new_p)
    para = Paragraph(new_p, anchor_p._parent)
    para.style = style
    para.add_run(text)
    if align:
        para.paragraph_format.alignment = align
    return para

# ── full content ───────────────────────────────────────────────────────

ABSTRACT = [
    ("前馈式三维重建模型（DUSt3R、MASt3R、MonST3R 等）近年来发展较快，但"
     "不同模型的运行环境、输入格式和输出内容差别很大。同时管理多个模型的"
     "实验时，任务怎么建、远端怎么跑、结果怎么收、记录怎么存，缺少统一的"
     "工具支持，实验过程难以追溯和比较。"),

    ("本文设计并实现了 3R All-in-One 多模型三维重建实验管理平台。系统分为"
     "本地桌面端（React + Tauri 2）、本地 FastAPI 后端和远端 GPU 服务器三层，"
     "通过 SSH/SCP 把模型推理交给远端完成，本地只负责任务状态、模型信息和"
     "结果索引。平台目前注册了 DUSt3R、MASt3R、MonST3R、Spann3R、Fast3R、"
     "Align3R、CUT3R 和 Dream3R 共 8 个模型，提供任务创建、参数配置、远端"
     "执行、日志回传、输出合同校验、实验记录包导出和 Agent 环境编排等功能。"
     "Dream3R v1.1 按候选几何融合模型接入，支持 synthetic 和 proposal-cache "
     "两种输入模式。"),

    ("测试方面，8 个模型蓝图全部通过有效性校验，174 条后端与 Agent 单元"
     "测试通过，前端构建正常，已打出 v0.5.0 Windows 安装包。演示实验用"
     " ETH3D delivery_area 场景的 12 张图像走完了从建任务到查结果的全流程。"),

    "关键词：三维重建；前馈式模型；实验管理平台；远端执行；Dream3R",
]

CH1_PARAS = {
    # heading 背景 (p46) → body p47, p48, p49
    47: ("三维重建是计算机视觉中的核心问题之一，目标是从图像、视频或多视角观"
         "测中恢复场景的三维结构，在机器人导航、增强现实、数字内容制作和空间"
         "理解等领域有实际用途。传统方法通常包括特征匹配、相机位姿估计、稀疏"
         "重建和稠密匹配等步骤，以 COLMAP 为代表的工具链已较为成熟，但对图像"
         "质量、视角重叠和相机运动稳定性要求较高。"),
    48: ("近年来，以 DUSt3R 为代表的前馈式三维重建模型快速发展。这类方法将"
         "部分传统几何步骤内化为网络前向推理，直接从输入图像预测点图或几何信"
         "息，降低了对相机标定和传统流水线的依赖。后续的 MASt3R、MonST3R、"
         "Spann3R、Fast3R、CUT3R 等工作分别在图像匹配、动态场景、多图输入和"
         "在线重建等方向进行了扩展。"),
    49: ("这些模型丰富了三维重建的可选方法，但同时引入了工程管理层面的困难。"
         "各模型的仓库结构、conda 环境、权重路径、运行命令和输出格式互不相同，"
         "手工逐一运行和记录时容易出现输入混淆、日志丢失、输出目录难以追溯"
         "等问题。"),

    # 研究现状 p51, p52, p53
    51: ("现有三维重建工具大致分为两类：一类是 COLMAP 等几何工具，流程完整"
         "但主要面向重建本身；另一类是各深度学习模型附带的推理脚本，便于复现"
         "论文示例，但缺少多模型在统一记录格式下的运行和对比能力。"),
    52: ("本项目要做的就是一个管理这些实验的本地软件。具体来说需要解决几个"
         "问题：模型信息怎么统一登记、任务怎么创建和派发、远端怎么运行和追"
         "踪、输出文件怎么回传和校验、实验参数和结果怎么保存下来方便以后再"
         "查。"),
    53: ("从结题角度看，单放几张效果图说明不了工作量。一个能建任务、跑远端、"
         "收结果、查日志的软件，才能直观说明项目做了什么。"),

    # 研究目标 p55, p56
    55: ("本项目的研究内容包括四个方面：（1）整理前馈式三维重建模型的输入输"
         "出规范，形成平台可读取的模型蓝图配置；（2）实现本地桌面软件与"
         " FastAPI 后端，完成模型注册、任务创建和状态管理；（3）打通基于"
         " SSH/SCP 的远端 GPU 运行流程，实现输入上传、命令执行、日志保存和结"
         "果回传；（4）接入本项目组自研的 Dream3R v1.1 候选几何融合模型，并完成"
         " Windows 稳定版打包。其中前三项是平台主体，第四项把一个自研模型也纳入"
         "同一套管理流程，用来验证平台对新模型的接入能力。"),
    56: ("项目的目标是做出一个能用的平台原型，把多模型实验的管理流程跑通，"
         "不追求完整复现每篇论文的 benchmark。从实际开发过程看，系统是一步步"
         "长出来的：v0.1 只能用 DUSt3R "
         "跑单模型 SSH 派发；v0.2 加了模型注册表和 AI 评估层；v0.3 把六个模型"
         "接进来并做了样例矩阵；v0.4 补了任务调度器和评估指标引擎；v0.5 加了"
         "对比可视化、参数模板和多服务器管理；结题阶段又接入 Dream3R 并补上了"
         " scene_meta 归一化和输出合同校验。"),
}

CH2_PARAS = {
    # 三维重建流程 p59, p60
    59: ("传统三维重建通常从多视图几何出发。SfM 方法通过特征匹配估计相机位"
         "姿和稀疏点云，MVS 方法在已知位姿基础上恢复稠密深度。这类方法的优"
         "点是几何约束清楚、结果可解释，但对弱纹理、重复纹理、运动模糊和宽"
         "基线图像较为敏感。"),
    60: ("从实验管理的角度，一次三维重建实验不只包含最终点云或深度图，还涉"
         "及输入图像、模型版本、运行参数、相机信息、中间日志和结果索引。缺少"
         "这些记录时，即使有可视化产物，实验过程也难以复查。"),

    # 前馈式模型 p62, p63
    62: ("前馈式三维重建模型将部分传统几何步骤内化为网络推理。DUSt3R 可在未"
         "标定图像条件下预测点图；MASt3R 强化了图像匹配与三维几何的联系；"
         "MonST3R 面向动态场景；Spann3R 和 CUT3R 关注序列状态管理；Fast3R "
         "则针对多图输入效率进行优化。"),
    63: ("这些模型的输入可能是图像对、多视图图像、视频片段或缓存文件，输出"
         "可能是点图、深度、mask、相机参数、可视化图片或 numpy 文件，平台层面"
         "无法强制统一。本项目采用统一任务记录、允许产物差异的策略，以"
         " scene_meta.json 作为各模型结果的统一索引格式。"),

    # 软件平台 p65, p66
    65: ("系统分成三层：本地桌面端、本地后端和远端 GPU 服务器。这样分的好处"
         "是本地不用装那些复杂的模型环境，只管任务状态、结果索引和用户交互，"
         "模型推理全交给远端机器跑。"),
    66: ("运行过程中平台会记几类关键信息：任务属于哪个模型、输入有哪些文件、"
         "跑的什么参数、当前状态走到哪一步、远端日志写了什么、输出目录里该有"
         "的文件到齐没有。这些信息都在，后面要复现或者比较才有依据。"),

    # Dream3R p68, p69
    68: ("Dream3R 是本项目组自行开发的模型，在平台里当做候选几何融合模型来"
         "用。它的标准输入是 proposal-cache（上游模型或缓存文件产出的候选几何"
         "结果），不是普通图片序列。为了演示时不依赖大文件，平台同时保留了"
         " synthetic 模式，可以直接走通演示路径。"),
    69: ("接入 Dream3R 主要是让平台能识别它的输入类型、参数、运行方式和输出"
         "产物，然后按跟其他模型一样的流程管理。当前版本任务完成后输出的记录"
         "包括任务摘要、dream3r_report.json 和可选的融合产物，在任务详情页"
         "统一查看。"),
}

CH3_PARAS = {
    # 总体设计 p72, p73
    72: ("系统分本地和远端两块：本地软件管任务——模型选择、建任务、看队列、"
         "查结果、改配置；本地后端存任务状态、模型注册信息和结果索引；远端"
         "服务器放模型权重和环境，负责真正跑推理。"),
    73: ("图 1　系统三层架构：桌面端 / 本地后端 / 远端 GPU 服务器（此处插入架构图）",),

    # 任务流程 p75, p76
    75: ("建任务时，平台会根据选的模型限制能选什么输入和参数：静态重建模型"
         "收图像，动态模型收视频或序列，Dream3R 可以选 synthetic 或者传"
         " proposal-cache 文件。提交之后后端把输入和配置传到远端，对应的"
         " runner 启动模型开始跑。"),
    76: ("跑的过程中平台持续记录状态变化和远端日志。任务跑成功了，产物经输"
         "出合同校验后进入结果索引；跑失败了，失败在哪一步、报什么错也会存"
         "下来，后续可以定位问题或者直接重试。"),

    # 模型接入 p78, p79
    78: ("新模型接入拆成四块：（1）agent/model_specs/ 下写一份 YAML 蓝图，描述"
         "环境、权重、构建步骤和 smoke 测试；（2）在 backend/model_registry.py "
         "里加一条模型记录，前端据此生成选择项和参数表单；（3）写一个"
         " runners/<model>_runner.py 负责封装具体的推理命令；（4）在"
         " backend/model_contracts.py 里定义输出合同，写清楚任务完成后必须有"
         "哪些文件。"),
    79: ("这样拆之后，接新模型不用动主逻辑，补齐上面四样东西就行。结题阶段"
         "的 8 个模型都是这么接的。"),

    # 软件使用 p81, p82
    81: ("实际用的时候大概是这样：打开软件先看队列里有什么任务，要跑新实验"
         "就点新建任务，选模型、配参数、传输入文件，提交后回队列等着，跑完了"
         "进详情页看产物、日志和合同校验结果。"),
    82: ("图 5　新建任务页：模型选择与参数配置（此处插入截图）"),

    # Dream3R实现 p84, p85
    84: ("Dream3R v1.1 按候选几何融合流程接入。它有 demo_mode（synthetic/cache）"
         "、domain、batch、views、patches、d_memory 和 device 这些参数，在"
         "建任务页面会根据模式不同显示不同的输入入口——synthetic 模式直接跑"
         "内置演示，cache 模式要上传 .pt 或 .pth 的 proposal-cache 文件。"),
    85: ("任务跑完后，Dream3R 的输出会整理成平台统一格式的记录：任务摘要加"
         " dream3r_report.json，再加可选的融合产物，在任务详情页和其他模型一样"
         "的方式查看。当前版本主要验证的是这条演示通路能不能走通——建任务、"
         "调 runner、收结果、在软件里查产物。"),
}

CH4_PARAS = {
    # 软件完成情况 p88, p89
    88: ("桌面端做完了模型选择、建任务、队列管理、结果检查和系统配置这些页面；"
         "后端实现了任务生命周期管理、模型注册、远端 SSH/SCP 调度和结果索引；"
         "远端服务器负责跑模型和生成产物。软件整体能跑通。"),
    89: ("调试过程中碰到了一些实现层面的问题。比较典型的一个：批量实验编排"
         "里的 run_experiment_from_template 最开始调了一个根本不存在的"
         " create_job(files=...) 方法，任务创建直接 TypeError，输入没存、任务"
         "也没派出去。修的时候改成按参数网格逐个建任务，用 save_inputs 存输入"
         "再返回真实 job_id，批量编排才真正能用。另外辅助评估配置页曾经因为前"
         "后端字段名对不上直接白屏，查了一圈发现是命名不一致，统一之后恢复。"),

    # 模型接入结果 p91
    91: ("平台注册了 8 个模型，各模型的接入状态见表 1。每个模型验证做到哪一"
         "步就写到哪一步，没有强行统一成同一级别。"),

    # 4.4 结果分析 p93
    93: ("总的来看，开题时设定的几个核心问题都解决了：模型信息统一记在注册"
         "表和蓝图里，任务通过软件建和派，结果跑完后有合同校验，参数和日志"
         "都存下来可以回头查。用 ETH3D delivery_area 的 12 张图走了一遍从建"
         "任务到查结果的完整流程，没有中断。定量指标方面，有 ground truth 时"
         "能算深度误差和点云密度，没有 ground truth 时记运行时间、完整性和日"
         "志；更全面的定量比较需要固定数据集，留给后续做。"),
}

CH5_PARAS = [
    ("本项目完成了一个面向前馈式三维重建模型的本地实验管理平台。当前版本"
     "已经把模型注册、任务创建、远端执行、日志回传、产物校验、实验记录和"
     "桌面打包这些环节跑通，用户可以在同一个软件里完成从建任务到查结果的"
     "完整过程。"),

    ("模型接入方面，平台注册了 8 个模型，各模型按实际验证深度标注状态。其中"
     " Dream3R 是本项目组自研的候选几何融合模型，能和 7 个开源模型走同一套接入"
     "和管理流程，说明平台不只适配现成模型，也能容纳自研模型。软件交付方面，"
     "系统通过了 8/8 蓝图校验和 174 条单元测试，前端构建正常，已打出 v0.5.0 "
     "Windows 安装包。后续如果继续做，主要方向是在固定数据集上补充定量对比、"
     "完善结果展示页面、扩展更多模型接入。"),
]

# ── apply replacements ────────────────────────────────────────────────

def apply(idx, new_text):
    if isinstance(new_text, tuple):
        # centred paragraph (figure caption)
        set_text(P[idx], new_text[0])
        P[idx].paragraph_format.alignment = CEN
    else:
        set_text(P[idx], new_text)

# abstract paragraphs 41-44
for i, txt in zip([41, 42, 43, 44], ABSTRACT):
    apply(i, txt)

for idx, txt in CH1_PARAS.items():
    apply(idx, txt)
for idx, txt in CH2_PARAS.items():
    apply(idx, txt)
for idx, txt in CH3_PARAS.items():
    apply(idx, txt)
for idx, txt in CH4_PARAS.items():
    apply(idx, txt)

# ch5 p95, p96, p97 → collapse to 2 paras + keep p97 blank
apply(95, CH5_PARAS[0])
apply(96, CH5_PARAS[1])
apply(97, "")   # was 3rd para, clear it

# ── figure captions already in CH3, but add missing ones as inserts ──
# 3.2 after p76: 图2 状态图, 图3 队列截图
cur = P[76]
for txt in [
    "图 2　任务生命周期状态流转（此处插入状态图）",
    "图 3　工作队列页：任务状态、筛选与失败原因摘要（此处插入截图）",
]:
    cur = insert_after(cur, txt, STYLE_BODY, CEN)

# 3.3 after p79: 图4
cur = insert_after(P[79], "图 4　新模型接入的四个稳定接口（此处插入流程图）", STYLE_BODY, CEN)

# 3.4 after p81: 图5 already set in p82 via apply; add 图6 after that
# p82 now contains 图5 caption; insert 图6 after p82
cur = insert_after(P[82],
    "图 6　任务详情页：产物列表、日志与输出合同校验分区（此处插入截图）",
    STYLE_BODY, CEN)

# 3.5 after p85: 图7
insert_after(P[85], "图 7　Dream3R 参数配置页（此处插入截图）", STYLE_BODY, CEN)

# 4.3 section: insert heading + body + figure captions before p92 (结果分析 H2)
# P[92] is the "结果分析" heading; 4.3 must precede it
p92_heading = P[92]
figures_43 = [
    ("图 8　演示输入：ETH3D delivery_area 12 帧（此处插入 contact sheet）", True),
    ("图 9　一次完整任务的输出产物与运行日志（此处插入截图）", True),
    ("· Agent 蓝图校验：8/8 valid", False),
    ("· Python 后端与 Agent 测试：174 条通过（另有 1 条按设计跳过）", False),
    ("· 前端构建：通过", False),
    ("· 稳定版产物：3R All-in-One_0.5.0_x64-setup.exe", False),
    ("图 10　发布检查 release_check 终端输出（此处插入截图）", True),
    ("图 11　Windows 稳定版安装包（此处插入安装包截图）", True),
    ("演示实验用 ETH3D delivery_area 场景的 12 张图跑了一次完整任务：在软件里"
     "选模型、传图、建任务，后端把输入传到远端，runner 跑模型并写日志，跑完后"
     "本地把结果目录同步回来，最后在任务详情页看输出和日志。", False),
]
# insert in reverse so each addprevious keeps order
new_43_head = OxmlElement('w:p')
p92_heading._p.addprevious(new_43_head)
head_para = Paragraph(new_43_head, p92_heading._parent)
head_para.style = STYLE_H2
head_para.add_run("演示样本与验证结果")

cur = head_para
for txt, centered in figures_43:
    cur = insert_after(cur, txt, STYLE_BODY, CEN if centered else None)

# 4.2 model table + table title (insert before p93 heading, after p91)
p91_new = P[91]
tbl_title = insert_after(p91_new, "表 1　各模型平台接入状态", STYLE_BODY, CEN)
MODELS = [
    ("模型", "平台定位", "接入状态"),
    ("DUSt3R", "静态重建基线", "baseline"),
    ("MASt3R", "静态匹配与重建", "smoke 验证"),
    ("MonST3R", "动态视频重建", "标准样本验证"),
    ("Spann3R", "空间记忆重建", "smoke 验证"),
    ("Fast3R", "多图快速重建", "smoke 验证（含 attention 回退）"),
    ("Align3R", "视频深度对齐", "runner 就绪，完整数据集验证待续"),
    ("CUT3R", "在线持久状态重建", "smoke 验证"),
    ("Dream3R", "候选几何融合", "v1.1 synthetic / cache 演示通路"),
]
tbl = doc.add_table(rows=len(MODELS), cols=3)
try:
    tbl.style = "Table Grid"
except Exception:
    pass
for r, row in enumerate(MODELS):
    for c, val in enumerate(row):
        cell = tbl.cell(r, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        if r == 0:
            run.bold = True
tbl_title._p.addnext(tbl._tbl)

# ── two new refs for Spann3R and CUT3R ────────────────────────────────
p107 = P[107]
cur = insert_after(p107,
    "Wang H., Agapito L. Spann3R: 3D Reconstruction with Spatial Memory[C]"
    "//International Conference on 3D Vision (3DV). 2025.",
    STYLE_REF)
insert_after(cur,
    "Wang Q., Zhang Y., Holynski A., Efros A. A., Kanazawa A. CUT3R: Continuous "
    "3D Perception Model with Persistent State[C]//Proceedings of the IEEE/CVF "
    "Conference on Computer Vision and Pattern Recognition (CVPR). 2025.",
    STYLE_REF)

doc.save(OUT)
print("saved:", OUT)
