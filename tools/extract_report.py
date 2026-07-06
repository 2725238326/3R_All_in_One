import sys
from docx import Document

src = sys.argv[1]
out = sys.argv[2]
doc = Document(src)
lines = []
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else ""
    lines.append(f"[{style}] {t}")

# also dump tables
for ti, tbl in enumerate(doc.tables):
    lines.append(f"\n=== TABLE {ti} ===")
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        lines.append(" | ".join(cells))

with open(out, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print(f"wrote {len(lines)} lines, {len(doc.tables)} tables")
