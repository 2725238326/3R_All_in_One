# -*- coding: utf-8 -*-
"""嵌入真实截图、删除无图占位、重排图号。

只有 4 个图有真实截图：
  工作队列页 -> 3.png   新建任务页 -> 5.png
  任务详情页 -> 6.png   Dream3R 参数页 -> 7.png
其余 7 个占位（架构/状态/接口/样本/输出/发布/安装包）删除。
保留图号重排为 图1~图4。
"""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

SRC = sys.argv[1]
OUT = sys.argv[2]
IMG_DIR = sys.argv[3].rstrip("/\\")

doc = Document(SRC)
CEN = WD_ALIGN_PARAGRAPH.CENTER

# 关键词 -> (图片文件, 新图号, 新标题, 图片宽度英寸)
KEEP = {
    "工作队列页": ("3.png", 1, "工作队列页：任务状态、筛选与失败原因摘要", 5.9),
    "新建任务页": ("5.png", 2, "新建任务页：模型选择与参数配置", 5.9),
    "任务详情页": ("6.png", 3, "任务详情页：产物列表、日志与输出合同校验分区", 5.9),
    "Dream3R 参数配置页": ("7.png", 4, "Dream3R 参数配置页", 3.3),
}
# 需要整段删除的占位关键词
DROP = ["系统三层架构", "任务生命周期", "新模型接入", "演示输入",
        "一次完整任务", "发布检查", "Windows 稳定版安装包"]


def remove_p(p):
    p._p.getparent().remove(p._p)


def insert_image_before(caption_p, img_path, width_in):
    new_p = OxmlElement('w:p')
    caption_p._p.addprevious(new_p)
    para = Paragraph(new_p, caption_p._parent)
    para.alignment = CEN
    para.add_run().add_picture(img_path, width=Inches(width_in))
    return para


def set_caption(p, text):
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    p.alignment = CEN


# 先快照所有段落（后续会增删，避免迭代时变动）
paras = list(doc.paragraphs)
removed, embedded = 0, 0
for p in paras:
    t = p.text.strip()
    if not t.startswith("图 "):
        continue
    # 删除无图占位
    if any(k in t for k in DROP):
        remove_p(p)
        removed += 1
        continue
    # 嵌入真图并重排图号
    for kw, (fname, num, title, w) in KEEP.items():
        if kw in t:
            insert_image_before(p, f"{IMG_DIR}/{fname}", w)
            set_caption(p, f"图 {num}　{title}")
            embedded += 1
            break

doc.save(OUT)
print(f"removed {removed} placeholders, embedded {embedded} images")
print("saved:", OUT)
