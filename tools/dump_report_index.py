import sys
from docx import Document

src = sys.argv[1]
doc = Document(src)
print(f"=== {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "?"
    txt = p.text.strip()
    short = (txt[:50] + "...") if len(txt) > 50 else txt
    print(f"{i:3d} [{style}] {short}")
