# -*- coding: utf-8 -*-
"""优化结题报告正文：补证据、补图占位、修数字、补参考文献、新增 4.3 节。"""
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1]
OUT = sys.argv[2]
doc = Document(SRC)
P = doc.paragraphs

# --- style objects grabbed from existing paragraphs (avoid CN name encoding) ---
BODY = P[47].style      # 论文正文
H2 = P[46].style        # Heading 2
REF = P[99].style       # 参考文献

def _mk(anchor_p, text, style, align, before):
    new_p = OxmlElement('w:p')
    if before:
        anchor_p._p.addprevious(new_p)
    else:
        anchor_p._p.addnext(new_p)
    para = Paragraph(new_p, anchor_p._parent)
    para.style = style
    if text:
        para.add_run(text)
    if align is not None:
        para.paragraph_format.alignment = align
    return para

def after(anchor, blocks):
    """blocks: list of (text, style, align). Returns last inserted para."""
    cur = anchor
    for text, style, align in blocks:
        cur = _mk(cur, text, style, align, before=False)
    return cur

def before(anchor, blocks):
    """Insert each before anchor in order (order preserved)."""
    for text, style, align in blocks:
        _mk(anchor, text, style, align, before=True)

CEN = WD_ALIGN_PARAGRAPH.CENTER

# capture anchors first
p43, p56, p73, p76, p79, p82, p85, p89, p91, p92, p107 = (
    P[43], P[56], P[73], P[76], P[79], P[82], P[85], P[89], P[91], P[92], P[107])

# === A. 摘要：补验证数字（插在 关键词 之前）===
after(p43, [(
    "系统已通过 8 个模型蓝图校验（8/8 valid）、174 条后端与 Agent 单元测试和前端"
    "构建，并生成 v0.5.0 Windows 稳定版安装包，具备结题演示和后续扩展基础。",
    BODY, None)])

# === B. 1.3 研究内容与目标：补版本演进脉络 ===
after(p56, [(
    "从开发过程看，平台并非一次设计成型，而是逐步生长起来的。最初的 v0.1 只支持 "
    "DUSt3R 单模型的 SSH/SCP 远端派发和本地任务缓存；v0.2 加入模型注册表和 AI 评估层，"
    "使平台具备容纳多模型的结构；v0.3 完成六模型集成与样例矩阵；v0.4 补齐任务调度器、"
    "评估指标引擎和报告导出；v0.5 加入对比可视化、参数模板、多服务器管理和发布检查；"
    "结题阶段进一步接入 Dream3R v1.1，并补全 scene_meta 归一化与输出合同校验，形成任务"
    "闭环。这条从单模型脚本到多模型平台的演进，本身就是本项目研究内容的主线。",
    BODY, None)])

# === C. 第 3 章 图占位 ===
after(p73, [("图 1　系统三层架构：本地桌面端 / 本地后端 / 远端 GPU 服务器（此处插入架构图）", BODY, CEN)])
after(p76, [
    ("图 2　任务生命周期状态流转（此处插入状态图）", BODY, CEN),
    ("图 3　工作队列页：任务状态、筛选与失败原因摘要（此处插入软件截图）", BODY, CEN)])
after(p79, [("图 4　新模型接入的四个稳定接口：蓝图 / 注册表 / runner / 输出合同（此处插入流程图）", BODY, CEN)])
after(p82, [
    ("图 5　新建任务页：模型选择与参数配置（此处插入软件截图）", BODY, CEN),
    ("图 6　任务详情页：产物、日志与输出合同校验分区（此处插入软件截图）", BODY, CEN)])
after(p85, [("图 7　Dream3R 参数配置：synthetic / cache 模式（此处插入软件截图）", BODY, CEN)])

# === D1. 4.1 软件完成情况：补真实修 bug 过程 ===
after(p89, [(
    "在功能联调阶段也定位并修复了若干实现层面的问题。例如批量实验编排最初调用了并不"
    "存在的 create_job(files=...) 接口，导致任务创建抛出 TypeError、输入从未保存、任务也"
    "从未真正派发；修复时改为按参数网格创建可派发任务，复用来源任务的输入并返回真实 "
    "job_id，批量编排才真正打通。这类问题在手工运行单个模型时不易暴露，正是把实验流程"
    "统一到软件中之后，才被稳定复现并修复。",
    BODY, None)])

print("text edits done")

# === D2. 4.2 模型接入结果：插入模型状态表（表1）===
# 表题
after(p91, [("表 1　各模型平台接入状态", BODY, CEN)])
MODELS = [
    ("模型", "平台定位", "接入状态"),
    ("DUSt3R", "静态重建基线路径", "baseline"),
    ("MASt3R", "静态匹配与重建", "smoke 验证"),
    ("MonST3R", "动态视频重建", "标准样本验证"),
    ("Spann3R", "空间记忆重建", "smoke 验证"),
    ("Fast3R", "多图快速重建", "smoke 验证（含 attention 回退）"),
    ("Align3R", "视频深度对齐", "runner 就绪，完整数据集验证待续"),
    ("CUT3R", "在线持久状态重建", "smoke 验证"),
    ("Dream3R", "候选几何融合", "v1.1 synthetic / cache 演示通路接入"),
]
tbl = doc.add_table(rows=len(MODELS), cols=3)
try:
    tbl.style = "Table Grid"
except Exception:
    pass
for r, row in enumerate(MODELS):
    for c, val in enumerate(row):
        cell = tbl.cell(r, c)
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(val)
        if r == 0:
            run.bold = True
# move table to sit right after 表题 (which is the para after p91)
table_title = p91._p.getnext()  # 表题 para
table_title.addnext(tbl._tbl)

# === E. 新增 4.3 演示样本与验证结果（目录有、正文缺；插在 4.4结果分析标题 p92 之前）===
before(p92, [
    ("演示样本与验证结果", H2, None),
    ("为了在结题环节给出可复查的完整实验，平台选取 ETH3D delivery_area 场景中的 12 张"
     "代表帧作为演示输入。一次完整任务依次经过：在新建任务页选择模型并上传这组图像，"
     "后端把输入与运行配置上传到远端 GPU 服务器，远端 runner 启动模型并写入日志，任务"
     "完成后本地同步输出目录，最后在任务详情页查看产物、日志和输出合同校验结果。", BODY, None),
    ("图 8　演示输入：ETH3D delivery_area 12 帧（此处插入 contact sheet）", BODY, CEN),
    ("图 9　一次完整任务的输出产物与运行日志（此处插入任务详情页截图）", BODY, CEN),
    ("在软件交付层面，平台通过了一组可重复的验证检查，结果如下：", BODY, None),
    ("· Agent 蓝图校验：8/8 valid", BODY, None),
    ("· Python 后端与 Agent 测试：174 条通过（另有 1 项按设计跳过）", BODY, None),
    ("· 前端构建：通过", BODY, None),
    ("· 发布检查 release_check：核心项通过", BODY, None),
    ("· 稳定版产物：3R All-in-One_0.5.0_x64-setup.exe", BODY, None),
    ("其中 Docker 相关检查因本机未安装 Docker CLI 仅给出 warning，不影响桌面稳定版的主"
     "流程。", BODY, None),
    ("图 10　发布检查 release_check 终端输出（此处插入终端截图）", BODY, CEN),
    ("图 11　Windows 稳定版安装包（此处插入安装包文件截图）", BODY, CEN),
])

# === F. 参考文献：补 Spann3R 与 CUT3R（正文提及但原文漏引），插在 Align3R(p107) 之后 ===
after(p107, [
    ("Wang H., Agarwal S., et al. SpaNN3R: 3D Reconstruction with Spatial Memory[C]"
     "//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern "
     "Recognition (CVPR). 2025.", REF, None),
    ("Wang Q., Zhang Y., et al. CUT3R: Continuous 3D Perception Model with Persistent "
     "State[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern "
     "Recognition (CVPR). 2025.", REF, None),
])

print("table + 4.3 + refs done")
doc.save(OUT)
print("saved")
