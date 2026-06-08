from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.table import Table
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"E:\Demo\3R_All_in_One")
TEMPLATE = ROOT / "开题报告-3R_all in one.docx"
OUT = ROOT / "开题报告-3R_all in one_完成版.docx"
PLACEHOLDER_DIR = ROOT / "docx_placeholders"


TOC_LINES = [
    "1 论文选题的意义及背景\t1",
    "1.1 研究背景\t1",
    "1.2 本课题的核心价值意义\t2",
    "2 国内外研究现状分析\t3",
    "2.1 基于几何的传统三维重建方法\t3",
    "2.2 基于学习的深度估计与多视图重建方法\t4",
    "2.3 前馈式三维重建模型\t5",
    "2.4 实验管理工具与平台需求\t6",
    "3 前期准备、研究目标与研究内容\t7",
    "3.1 前期准备与已有工作\t7",
    "3.2 研究目标\t8",
    "3.3 研究内容\t9",
    "3.4 拟解决的关键问题\t10",
    "3.5 初拟论文提纲\t10",
    "4 研究方案及可行性分析\t11",
    "4.1 平台总体技术路线\t11",
    "4.2 核心模块设计\t12",
    "4.3 实验设计\t13",
    "4.4 可行性分析\t14",
    "5 预期达到的目标和研究成果\t15",
    "6 论文工作计划\t16",
    "7 参考文献\t17",
]


BODY = [
    ("h1", "论文选题的意义及背景"),
    ("h2", "研究背景"),
    ("p", "三维重建是计算机视觉中的基础问题，目标是根据图像、视频或多视图观测恢复场景的三维结构。该方向在机器人导航、增强现实、数字内容制作、空间测量和三维场景理解等领域都有实际用途。传统三维重建流程通常由特征提取与匹配、相机位姿估计、稀疏重建、稠密匹配和点云融合等步骤组成，典型工具包括 COLMAP 等。这类方法在纹理充足、视角变化适中、相机运动稳定的场景中表现较好，但在弱纹理、重复纹理、动态物体、宽基线图像和未标定输入下，容易出现匹配失败、位姿漂移或重建不完整的问题。"),
    ("p", "近几年，前馈式三维重建模型逐渐受到关注。以 DUSt3R 为代表的方法尝试直接从图像中预测点图、深度或相机相关几何表示，降低了对传统几何流程和相机标定信息的依赖。后续工作又在图像匹配、动态场景、多视图输入、长序列处理和在线重建等方向继续扩展，例如 MASt3R、MonST3R、Spann3R、Fast3R、CUT3R 等。这些模型为三维重建实验提供了新的工具，也使不同输入条件下的快速推理和结果展示变得更加可行。"),
    ("p", "然而，从实际学习和实验过程来看，多模型比较的成本并不只来自算法本身，还大量来自工程使用环节。研究者需要分别配置 conda 环境、准备权重、修改脚本参数、上传输入数据、运行远端 GPU 命令、下载结果并整理日志。不同模型的输出可能是点图、深度、mask、相机参数或可视化文件，文件组织方式也不一致。如果只保留最终截图或个别成功样例，很难判断某个模型在固定输入上的真实表现，也不方便记录失败原因。"),
    ("p", "因此，本课题拟设计一个面向前馈式三维重建模型的聚合管理平台。课题重点不在重新提出一种三维重建网络，而是放在模型实验流程的组织、运行和记录上，尝试建立统一的模型描述、任务提交、远端运行、结果回传和指标评估方式。该平台可以作为后续模型复现、横向比较和项目展示的基础工具，也有助于把实验失败、环境问题和输出差异记录下来，减少重复调试。"),
    ("h2", "本课题的核心价值意义"),
    ("p", "基于上述背景可以看出，前馈式三维重建模型的发展速度较快，但模型之间的工程接口并不统一。对学生实验和项目开发而言，真正影响效率的往往不是单次推理，而是环境配置、输入准备、运行脚本、结果查找和失败复盘这些重复环节。本课题希望把这些环节整理成相对统一的平台流程，使模型复现和对比不再完全依赖零散脚本和人工记录。"),
    ("p", "本课题的意义主要体现在三个方面。第一，通过模型注册和配置文件，把不同模型的输入、输出、权重、运行命令和资源需求统一记录，便于后续扩展。第二，通过任务提交、远端执行和结果同步流程，降低本地设备运行大模型的压力，提高实验过程的可追溯性。第三，通过结果元数据和指标表，把成功样例、失败样例、日志和可视化产物一起保存，为后续中期检查、论文写作和项目展示提供依据。"),

    ("h1", "国内外研究现状分析"),
    ("p", "围绕三维重建，国内外研究已经形成了较为丰富的技术体系。从传统几何方法到深度学习方法，再到近年来的前馈式重建模型，各类方法在输入条件、输出形式和工程使用方式上差异明显。本课题关注的是前馈式三维重建模型的聚合管理，因此既需要了解三维重建方法本身，也需要分析多模型实验组织工具的现实需求。"),
    ("h2", "基于几何的传统三维重建方法"),
    ("p", "传统三维重建主要依赖多视图几何。SfM 方法通过图像特征匹配估计相机位姿和稀疏点云，MVS 方法在已知相机位姿的基础上恢复稠密深度或点云。COLMAP 等工具把这些步骤组织成较成熟的流程，在静态场景和高质量图像中有较强可用性。这类方法的优势是几何约束清楚、结果可解释，且有较长时间的工程积累。"),
    ("p", "其局限主要体现在输入条件要求较高。若图像之间重叠区域不足、纹理信息弱、物体发生运动或相机轨迹不稳定，特征匹配和位姿估计容易失败。对于初学者或多模型实验场景，传统流程还需要较多参数设置和中间结果判断，不利于快速比较不同方法的实际效果。"),
    ("h2", "基于学习的深度估计与多视图重建方法"),
    ("p", "深度学习方法在深度估计、多视图立体和场景重建中取得了较多进展。MVSNet 等方法利用神经网络构建立体匹配代价体，提升了深度估计的稳定性；NeuralRecon 等方法尝试结合 RGB-D 或序列信息进行场景级重建；NeRF 及其后续方法则从隐式表示角度建模场景，可以生成连续视角下的渲染结果。"),
    ("p", "这类方法在固定数据分布或特定任务上效果较好，但也存在使用条件限制。例如部分方法需要相机内外参、深度监督或较长优化时间；部分方法面向单一场景训练，难以直接用于快速批量比较。对于本课题关注的多模型实验管理，这些方法说明了评价指标、数据集选择和运行记录的重要性，也提示平台需要保存相机参数、深度、点云和运行耗时等多类信息。"),
    ("h2", "前馈式三维重建模型"),
    ("p", "前馈式三维重建模型试图减少传统流程中的显式特征匹配和手工几何步骤。DUSt3R 将输入图像映射到统一的三维表示，能够在未标定图像对或多视图条件下直接输出点图。MASt3R 在图像匹配和三维几何约束之间建立联系，适合整理图像匹配与重建任务之间的关系。MonST3R 面向动态场景和视频输入，尝试处理物体运动带来的几何变化。Spann3R、Fast3R 和 CUT3R 等工作分别关注长序列、多图输入、运行效率和在线重建等问题。"),
    ("p", "这些方法的共同特点是输入和输出更接近端到端模型推理，使用时通常只需要准备图像、视频或少量参数即可得到几何结果。但它们的使用难点也比较明显：每个模型仓库的依赖版本不同，有些模型需要编译 CUDA 扩展或特殊算子，有些模型权重较大，部分模型对显存和输入长度有明确限制。模型之间的输出文件也不统一，直接比较会遇到格式转换和指标口径问题。"),
    ("fig", "图 2.1 前馈式三维重建模型实验流程示意图（占位）"),
    ("h2", "实验管理工具与平台需求"),
    ("p", "现有开源工具多数服务于单一算法或传统重建流程。COLMAP 有较完整的命令行和图形界面，但主要面向几何重建。深度学习模型通常提供独立的推理脚本和 demo，方便复现论文示例，但较少考虑多个前馈式三维重建模型在同一输入和同一记录格式下的比较。实际项目中常见做法是使用脚本手动运行，再用表格或文件夹记录结果。"),
    ("p", "本课题计划补充的是实验组织层面的工具。平台将模型差异放在执行器和配置文件中处理，把任务、日志、产物和指标保存在统一位置。这样可以让后续实验更容易追溯，也能帮助使用者分辨模型本身能力、运行环境问题和输入数据问题之间的区别。"),

    ("h1", "前期准备、研究目标与研究内容"),
    ("h2", "前期准备与已有工作"),
    ("p", "开题阶段已经完成了相关模型族的初步阅读，重点整理各方法的输入形式、输出表示、适用场景和可能失败条件。已阅读和整理的代表模型包括 DUSt3R、MASt3R、MonST3R、Spann3R、Fast3R、CUT3R 和 Align3R 等。整理内容包括模型是否需要相机参数，主要输入是图像对、多视图图像还是视频序列，输出中是否包含点图、深度、mask、相机信息或可视化结果。"),
    ("p", "我对模型失败条件也进行了初步整理。静态图像对模型可能在低纹理、视角差过大、重叠区域不足或强反光场景中表现不稳定；动态场景模型可能受到运动模糊、遮挡和长序列显存限制影响；部分模型在 Windows 本地直接运行困难，需要远端 Linux GPU 环境；一些模型需要编译 CUDA 扩展，编译失败会导致推理脚本无法启动。平台设计中需要把这些失败情况记录为日志和状态，成功结果和失败原因都应保留下来。"),
    ("p", "在平台方案上，已经形成了基本框架：桌面前端用于选择模型、上传输入和查看结果；本地后端用于保存任务、校验配置和提供 API；远端调度用于通过 SSH/SCP 上传输入、启动命令和同步产物；模型执行器用于封装具体模型的运行脚本。该结构可以把平台通用逻辑和模型专有命令分开，后续新增模型时主要补充模型配置和执行器。"),
    ("p", "在接口和文件格式方面，开题阶段拟定了三个关键状态文件：job.json 记录模型、输入、参数和输出目录；status.json 记录运行状态、进度和错误信息；scene_meta.json 记录结果产物、相机信息和场景元数据。后续实验中，这些文件将作为任务复现和结果对比的依据。"),
    ("h2", "研究目标"),
    ("p", "本课题计划完成一个可运行的前馈式三维重建模型实验管理原型。目标尽量对应可检查的产出和实验。"),
    ("p", "1、建立模型注册与描述方式。产出包括模型配置文件、字段说明和校验脚本，后续可以检查每个模型是否记录了环境、输入、输出、权重、运行命令和资源需求。"),
    ("p", "2、实现任务提交与远端运行流程。产出包括任务 API、SSH/SCP 调用逻辑和 runner 脚本，后续可以用固定输入检查任务是否能生成配置、上传、运行、回传和保存日志。"),
    ("p", "3、建立结果解析与产物管理方式。产出包括 scene_meta.json 规范、结果目录结构和可视化入口，后续可以检查不同模型输出是否能被平台索引。"),
    ("p", "4、设计多模型对比实验。产出包括样本清单、指标表、对照模型列表和失败记录表，后续可以在固定样本上比较模型运行状态、耗时和输出质量。"),
    ("p", "5、形成项目文档和开题到中期材料。产出包括接口说明、实验设计说明、运行记录和阶段报告，方便复现实验过程。"),
    ("h2", "研究内容"),
    ("h3", "模型注册与配置模块"),
    ("p", "该模块输入为不同 3R 模型的论文信息、代码仓库信息、运行环境、权重路径、输入要求和输出格式。处理方式是将这些信息整理为统一的模型配置文件，并提供字段校验。输出为平台可读取的模型注册表。验证方式包括检查每个模型配置是否包含必要字段，以及平台能否按模型类型筛选和展示模型信息。"),
    ("h3", "任务提交与远端执行模块"),
    ("p", "该模块输入为用户选择的模型、输入图像或视频、参数配置和远端服务器信息。平台后端生成 job.json，通过 SSH/SCP 将输入和配置上传到远端目录，再调用对应 runner 执行模型推理。运行过程中记录日志和 status.json，完成后将输出目录同步回本地。该模块重点解决任务从本地到远端的完整闭环问题。"),
    ("h3", "结果解析与展示模块"),
    ("p", "该模块输入为模型运行产生的点图、深度、mask、相机信息、图片或视频产物。处理方式是由 runner 或后端整理 scene_meta.json，记录产物路径、文件类型、场景名称和可选相机信息。平台前端根据元数据展示结果，使不同模型的输出能够被归到同一任务目录下。"),
    ("h3", "指标评估与对比模块"),
    ("p", "该模块输入为固定样本的多模型输出和可选真值数据。若数据集提供深度、相机位姿或点云真值，计划计算 AbsRel、RMSE、ATE、RPE、Chamfer Distance、F-score 等指标。若样本没有真值，则记录运行耗时、显存需求、成功率、输出完整性、可视化质量和失败原因。"),
    ("h3", "文档与实验记录模块"),
    ("p", "该模块输入为模型配置、任务记录、测试输出和实验结果。处理方式是整理接口文档、运行说明、模型接入说明和实验记录模板。输出为项目文档、阶段报告和后续论文材料。该模块看似偏工程文档，但对后续复现实验和说明平台价值非常重要。"),
    ("h2", "拟解决的关键问题"),
    ("p", "1、不同模型的运行环境、输入格式和输出结构差异较大，平台需要在不强行改动模型仓库的前提下完成统一接入。"),
    ("p", "2、远端 GPU 任务存在上传、运行、日志回传和异常中断等环节，平台需要保证每一步都有状态记录，避免任务失败后无法追溯。"),
    ("p", "3、不同模型产物不完全一致，部分指标无法直接计算，因此需要设计可选字段和失败原因记录，而不是简单要求所有模型输出完全相同。"),
    ("p", "4、实验结果需要兼顾定量指标和实际可用性，既记录有真值数据集上的误差，也记录无真值样例中的运行耗时、输出完整性和失败类型。"),
    ("h2", "初拟论文提纲"),
    ("p", "本课题的论文初拟提纲如下所示："),
    ("p", "1 绪论"),
    ("p", "1.1 课题背景及意义"),
    ("p", "1.2 国内外研究现状"),
    ("p", "1.3 课题研究目标与主要内容"),
    ("p", "2 前馈式三维重建模型与实验管理需求分析"),
    ("p", "2.1 代表模型输入输出特征"),
    ("p", "2.2 多模型运行与对比中的工程问题"),
    ("p", "3 聚合管理平台总体设计"),
    ("p", "3.1 系统架构设计"),
    ("p", "3.2 模型注册与任务管理模块"),
    ("p", "3.3 远端 runner 与结果同步模块"),
    ("p", "4 平台实现与实验验证"),
    ("p", "4.1 模型接入实验"),
    ("p", "4.2 多模型对比实验"),
    ("p", "4.3 失败案例与结果分析"),
    ("p", "5 总结与展望"),

    ("h1", "研究方案及可行性分析"),
    ("h2", "平台总体技术路线"),
    ("p", "本研究拟采用本地桌面端加远端 GPU 执行的总体方案。本地侧负责模型管理、任务创建、状态展示和结果查看；远端侧负责真正的模型推理和产物生成；二者之间通过 SSH/SCP 完成输入上传、命令启动和结果同步。这样设计的原因是前馈式 3R 模型通常依赖 GPU，且环境配置较重，本地电脑更适合承担管理和展示工作。"),
    ("p", "本课题的技术路线可以概括为：输入数据 -> 样本整理 -> 模型选择与参数配置 -> 生成 job.json -> 后端校验 -> SSH/SCP 上传 -> 远端 runner 执行 -> 生成 status.json 和日志 -> 输出点图、深度、mask 或相机信息 -> 生成 scene_meta.json -> 本地同步 -> 可视化与指标评估 -> 对比表格和失败记录。"),
    ("fig", "图 4.1 平台总体架构图（占位）"),
    ("p", "平台技术栈拟采用 Tauri 2 与 React 作为桌面前端，FastAPI 作为本地后端，Python runner 作为模型执行器，SSH/SCP 作为远端传输方式。前端负责输入选择、任务列表和结果查看；后端负责任务状态、模型注册、接口调用和文件保存；远端服务器负责实际模型推理。该路线便于先完成小规模原型，再逐步补充更多模型和指标。"),
    ("h2", "核心模块设计"),
    ("p", "模型注册模块采用配置文件保存模型基本信息，包括模型名称、论文或仓库地址、运行环境、权重路径、输入类型、输出类型、默认命令、显存需求和已知限制。平台启动时读取注册表，并在前端展示可用模型。这样可以避免把模型差异写死在平台主逻辑中。"),
    ("p", "任务管理模块负责生成 job.json 和任务目录。一次任务至少包含模型标识、输入路径、参数、输出目录、创建时间和运行状态。任务提交后，本地后端将输入和配置上传到远端目录，远端 runner 根据模型配置启动推理脚本。运行过程中持续写入 status.json 和日志，任务完成后再同步结果文件。"),
    ("p", "结果解析模块负责整理 scene_meta.json。该文件记录点图、深度图、mask、相机参数、可视化图片、视频和日志等产物路径。如果某个模型没有输出某类文件，则对应字段可以为空，同时记录原因。通过这种方式，平台既能统一索引结果，又不会强行要求所有模型具有完全一致的输出形式。"),
    ("fig", "图 4.2 任务提交与结果回传流程图（占位）"),
    ("h2", "实验设计"),
    ("h3", "数据集与样本选择"),
    ("p", "实验计划分为小样本功能验证和公开数据集对比两部分。小样本功能验证拟选取少量图像对、多视图图像和短视频片段，用于检查平台能否正常提交任务、运行模型和回传结果。这部分样本不追求覆盖所有场景，重点是让每类输入都能被平台正确处理。"),
    ("p", "公开数据集对比拟参考以下数据来源：静态多视图场景可参考 DTU、ETH3D、Tanks and Temples；室内 RGB-D 或相机轨迹场景可参考 ScanNet、7-Scenes；动态或视频输入可选取公开视频数据和自建短序列样例。具体数据集会根据模型支持的输入格式、下载成本和真值可用性进行筛选。"),
    ("h3", "对照模型"),
    ("p", "前期计划以 DUSt3R 作为基础前馈式模型进行流程验证，再逐步接入 MASt3R、MonST3R、Spann3R、Fast3R、CUT3R 等模型。若时间允许，可加入 COLMAP 作为传统几何方法参考，用于比较前馈式模型与传统流程在固定样本上的差异。对照实验不只记录最终数值，也记录模型是否能成功启动、是否需要额外环境修复、是否能输出平台需要的元数据。"),
    ("h3", "评价指标"),
    ("p", "实验指标按数据条件分为两类。有真值数据时，拟采用深度误差、相机轨迹误差和点云指标。例如 AbsRel、RMSE 用于深度或距离误差，ATE、RPE 用于相机轨迹，Chamfer Distance、F-score 用于点云几何差异。不同指标是否适用，需要根据模型输出和数据集真值决定。"),
    ("p", "无真值样本时，拟记录运行耗时、任务成功率、输出文件完整性、点云或深度图是否可视化、日志是否包含错误、显存或输入长度限制等信息。这类记录不能替代几何精度指标，但可以反映平台管理和模型使用中的实际问题。"),
    ("h3", "实验类型"),
    ("p", "计划开展四类实验：第一，功能验证实验，检查任务创建、上传、远端运行、日志保存和结果回传是否能完成；第二，多模型对比实验，在相同样本上运行多个模型，记录耗时、输出类型、指标和失败情况；第三，参数影响实验，对输入分辨率、帧数、参数档位等进行小范围调整，观察运行时间和输出变化；第四，失败案例分析，对失败任务按环境错误、输入不匹配、显存不足、模型推理失败、输出缺失等类别记录，形成后续修复依据。"),
    ("h2", "可行性分析"),
    ("p", "从技术基础看，本课题使用的 Tauri、React、FastAPI、Python 和 SSH/SCP 都是成熟工具，适合搭建本地桌面端和远端执行流程。前期已经初步整理了模型族、接口格式和 runner 设计，可以支持原型开发。远端 GPU 服务器可以承担模型推理任务，本地电脑主要进行任务管理和结果展示，计算压力相对较小。"),
    ("p", "从数据和模型来源看，DUSt3R、MASt3R 等模型已有公开论文和代码，部分数据集也提供公开样本和真值。开题阶段已经初步梳理了候选模型、候选数据集和候选指标，后续可以从小样本开始验证，再扩展到公开数据集。"),
    ("p", "主要风险有四类。第一，部分模型环境复杂，CUDA、PyTorch 或编译扩展可能不兼容；应对方式是为每个模型单独记录环境和已知问题，优先接入环境稳定的模型。第二，不同模型输出不一致，可能导致某些指标无法直接计算；应对方式是先定义产物索引和可选字段，不强制所有模型输出完全相同内容。第三，公开数据集真值和模型输出之间可能存在尺度、坐标系或格式差异；应对方式是先进行小规模检查，记录无法比较的原因。第四，时间有限，无法完成所有模型的完整实验；应对方式是优先保证一个基础模型和一个扩展模型跑通，再逐步增加模型数量。"),

    ("h1", "预期达到的目标和研究成果"),
    ("p", "本课题预期完成一个面向前馈式三维重建模型的实验管理原型，包括前端界面、本地后端和远端 runner。平台能够完成模型查询、输入选择、任务创建、远端运行、日志保存、结果同步和基本结果查看，为后续模型复现和实验对比提供统一入口。"),
    ("p", "在模型接入方面，预期形成若干模型配置文件，记录 DUSt3R、MASt3R、MonST3R、Spann3R、Fast3R、CUT3R 等模型的输入、输出、环境和运行方式。配置文件将作为模型注册表的一部分，支持后续继续扩展其他模型。"),
    ("p", "在实验记录方面，预期形成一套任务记录格式，包括 job.json、status.json 和 scene_meta.json；形成一组固定实验样本和样本说明，覆盖图像对、多视图图像和视频片段；形成基础指标脚本和对比表格，记录运行耗时、成功率、输出完整性和可计算的几何指标。"),
    ("p", "在文档和论文材料方面，预期完成项目说明文档、接口文档、模型接入说明、失败案例记录表、阶段报告和后续论文材料。通过这些成果，可以较清楚地展示平台的设计思路、实现过程和实验价值。"),

    ("h1", "论文工作计划"),
    ("p", "论文工作计划如表 6.1所示："),
    ("table_caption", "表 6.1毕业设计论文工作计划"),
    ("plan_table", ""),

    ("h1", "7参考文献"),
]


REFERENCES = [
    "Schönberger J. L., Frahm J.-M. Structure-from-Motion Revisited[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016.",
    "Schönberger J. L., Zheng E., Pollefeys M., Frahm J.-M. Pixelwise View Selection for Unstructured Multi-View Stereo[C]//European Conference on Computer Vision. 2016.",
    "Yao Y., Luo Z., Li S., Fang T., Quan L. MVSNet: Depth Inference for Unstructured Multi-view Stereo[C]//European Conference on Computer Vision. 2018.",
    "Sun J., Xie Y., Chen L., Zhou X., Bao H. NeuralRecon: Real-Time Coherent 3D Reconstruction from Monocular Video[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2021.",
    "Wang S., Leroy V., Cabon Y., Chidlovskii B., Revaud J. DUSt3R: Geometric 3D Vision Made Easy[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. 2024.",
    "Leroy V., Cabon Y., Revaud J. MASt3R: Grounding Image Matching in 3D[J/OL]. 2024.",
    "Zhang J., Herrmann C., Hur J., Jampani V., Darrell T., Cole F., Sun D., Yang M.-H. MonST3R: A Simple Approach for Estimating Geometry in the Presence of Motion[J/OL]. arXiv, 2024.",
    "Spann3R、Fast3R、CUT3R、Align3R 相关论文、代码仓库和模型说明资料.",
]


PLAN_ROWS = [
    ("1-2", "阅读 DUSt3R、MASt3R、MonST3R 等代表模型论文和代码说明，整理输入、输出、环境和失败条件。"),
    ("3-4", "完成平台模块设计，确定模型配置字段、任务状态文件和结果元数据格式。"),
    ("5", "开发本地后端基础 API，完成任务创建、配置校验和任务目录生成。"),
    ("6-7", "开发前端基础页面和远端 runner 原型，完成小样本任务提交与日志记录。"),
    ("8-9", "接入 2-3 个代表模型，准备固定样本，完成多模型输出保存和基础对比表。"),
    ("10", "补充结果解析、可视化入口和失败记录表，整理阶段性实验结果。"),
    ("11-12", "完善指标脚本、接口文档和中期材料，根据测试情况调整平台结构。"),
    ("后续", "根据中期反馈继续补充模型、改进结果展示，完善毕业论文和答辩材料。"),
]


def clear_para(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_para_text(paragraph, text: str):
    clear_para(paragraph)
    paragraph.add_run(text)


def make_placeholder(path: Path, label: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 560), "white")
    draw = ImageDraw.Draw(img)
    border = (120, 120, 120)
    draw.rectangle((18, 18, 1182, 542), outline=border, width=3)
    draw.line((18, 92, 1182, 92), fill=(210, 210, 210), width=2)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    try:
        title_font = ImageFont.truetype(str(font_path), 38)
        body_font = ImageFont.truetype(str(font_path), 30)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.text((60, 42), "图片占位", fill=(70, 70, 70), font=title_font)
    lines = [label, "后续替换为正式示意图或实验结果图"]
    y = 210
    for line in lines:
        box = draw.textbbox((0, 0), line, font=body_font)
        x = (1200 - (box[2] - box[0])) // 2
        draw.text((x, y), line, fill=(90, 90, 90), font=body_font)
        y += 56
    img.save(path)


def remove_body_from_first_heading(doc: Document):
    body = doc.element.body
    start_idx = None
    for idx, child in enumerate(list(body)):
        if child.tag.endswith("}p"):
            text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t"))
            if text == "论文选题的意义及背景":
                start_idx = idx
                break
    if start_idx is None:
        raise RuntimeError("Cannot find body start heading in template.")
    for child in list(body)[start_idx:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def replace_toc_with_static_lines(doc: Document):
    body = doc.element.body
    children = list(body)
    toc_title_idx = None
    section_break_idx = None
    for idx, child in enumerate(children):
        if not child.tag.endswith("}p"):
            continue
        text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t")).strip()
        if text == "目录":
            toc_title_idx = idx
            continue
        if toc_title_idx is not None and child.xpath(".//w:sectPr"):
            section_break_idx = idx
            break
    if toc_title_idx is None or section_break_idx is None:
        raise RuntimeError("Cannot locate TOC range in template.")

    section_break_para = None
    for para in doc.paragraphs:
        if para._p is children[section_break_idx]:
            section_break_para = para
            break
    if section_break_para is None:
        raise RuntimeError("Cannot bind TOC section break paragraph.")

    for child in children[toc_title_idx + 1 : section_break_idx]:
        body.remove(child)

    for line in TOC_LINES:
        head = line.split("\t", 1)[0]
        number = head.split(" ", 1)[0]
        level = number.count(".") + 1
        style = f"toc {min(level, 3)}"
        section_break_para.insert_paragraph_before(line, style=style)


def append_plan_table(doc: Document, template_table):
    tbl_xml = deepcopy(template_table._tbl)
    doc.element.body.insert(len(doc.element.body) - 1, tbl_xml)
    table = Table(tbl_xml, doc)
    for i, (week, task) in enumerate(PLAN_ROWS, start=1):
        table.cell(i, 0).text = week
        table.cell(i, 1).text = task
    return table


def build():
    doc = Document(str(TEMPLATE))
    plan_table_template = doc.tables[5]

    for para in doc.paragraphs:
        if para.text.strip() == "基于Pointmap的动态场景三维重建方法研究":
            set_para_text(para, "面向前馈式三维重建模型的聚合管理平台设计与实现")
            break

    replace_toc_with_static_lines(doc)

    remove_body_from_first_heading(doc)

    placeholder_idx = 1
    for kind, text in BODY:
        if kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(text, style="Heading 3")
        elif kind == "p":
            doc.add_paragraph(text, style="论文正文")
        elif kind == "fig":
            img_path = PLACEHOLDER_DIR / f"figure_placeholder_{placeholder_idx}.png"
            make_placeholder(img_path, text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Inches(5.7))
            doc.add_paragraph(text, style="图")
            placeholder_idx += 1
        elif kind == "table_caption":
            doc.add_paragraph(text, style="表题")
        elif kind == "plan_table":
            append_plan_table(doc, plan_table_template)
        else:
            raise ValueError(kind)

    for ref in REFERENCES:
        doc.add_paragraph(ref, style="参考文献")

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
